
from docx import Document
from pathlib import Path

import os
import pandas as pd

class CustomSurveyWordTable:
    '''自定义工勘Word表格'''
    def __init__(self, 
                 wordTemplatePath: str,
                 materialsExcelPath: str) -> None:
        """_summary_

        :param wordTemplatePath: word模板路径, 该模板是固定的
        :param materialsExcelPath: 材料表路径，该材料表是固定的
        """

        self.doc = Document(wordTemplatePath)
        self.tables = self.doc.tables

        if self.tables:
            print(f"文档中共有 {len(self.tables)} 个表格")
            self.table = self.tables[0]
            self.rows = len(self.table.rows)
            self.cols = len(self.table.columns)
            print(f"表格1: {self.rows} 行, {self.cols} 列")
        else:
            raise ValueError("文档中没有表格")
        
        self.df = pd.read_excel(materialsExcelPath, 
                            sheet_name = 'Sheet1',
                            usecols = 'P:X',  # 指定P到X列
                            skiprows = 2,     # 跳过前2行（读取从第3行开始）
                            nrows = 98)       # 读取98行（3-100行）
        
        
        
